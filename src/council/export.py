"""Export final paper to Markdown or Word (.docx), optionally with embedded images."""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Literal

from council.sections import article_body
from council.storage import SessionStore


def normalize_format(fmt: str) -> Literal["md", "docx"]:
    f = fmt.lower().strip().lstrip(".")
    if f in {"md", "markdown"}:
        return "md"
    if f in {"docx", "word", "doc"}:
        return "docx"
    raise ValueError(f"Unsupported format: {fmt}. Use md or docx/word.")


def resolve_final_markdown(store: SessionStore, *, with_images: bool = False) -> Path:
    candidates: list[Path] = []
    if with_images:
        candidates.append(store.path / "final" / "paper_with_figures.md")
    candidates.extend(
        [
            store.path / "final" / "paper_final.md",
            store.path / "draft" / "paper_v1.md",
        ]
    )
    for p in candidates:
        if p.exists() and p.stat().st_size > 0:
            text = p.read_text(encoding="utf-8", errors="replace")
            if text.strip().startswith("# FAILED"):
                continue
            return p
    raise FileNotFoundError(
        f"No final paper found in session {store.session_id}. "
        "Run finalize first (council resume <id> --from finalize)."
    )


def export_session(
    store: SessionStore,
    fmt: str = "md",
    out: Path | None = None,
    *,
    title: str | None = None,
    with_images: bool = False,
    ensure_images: bool = False,
    config: object | None = None,
) -> Path:
    """
    Export session paper to md or docx.

    with_images:
      - Prefer final/paper_with_figures.md
      - Resolve/convert figures to PNG and embed in Word
    ensure_images:
      - If with_images and figures missing, call generate_images (needs config)
    """
    fmt_n = normalize_format(fmt)

    if with_images:
        figs_md = store.path / "final" / "paper_with_figures.md"
        imgs_dir = store.path / "final" / "images"
        has_png = imgs_dir.exists() and any(imgs_dir.glob("figure_*.png"))
        if (not figs_md.exists() or not has_png) and ensure_images:
            if config is None:
                raise RuntimeError(
                    "with-images requires figures. Run `council images <id>` first, "
                    "or pass config so export can generate them."
                )
            # Lazy async run from sync context
            import asyncio

            from council.images import generate_images

            asyncio.run(generate_images(config, store))  # type: ignore[arg-type]
        elif not figs_md.exists():
            raise FileNotFoundError(
                "No paper_with_figures.md yet. Run:\n"
                f"  council images {store.session_id}\n"
                f"or: council word {store.session_id} --with-images --generate-images"
            )

    src = resolve_final_markdown(store, with_images=with_images)
    text = src.read_text(encoding="utf-8")
    # With images: strip revision metadata but keep the # Figures gallery.
    body = _article_body_keep_figures(text) if with_images else article_body(text)

    meta = store.load_meta()
    doc_title = title or meta.get("title") or "Paper"

    export_dir = store.path / "export"
    export_dir.mkdir(parents=True, exist_ok=True)

    if with_images:
        from council.media import rewrite_md_images_to_png

        body = rewrite_md_images_to_png(body, store.path)

    if fmt_n == "md":
        dest = out or (
            export_dir / ("paper_with_figures.md" if with_images else "paper_final.md")
        )
        dest = Path(dest)
        dest.parent.mkdir(parents=True, exist_ok=True)
        dest.write_text(body, encoding="utf-8")
        key = "export_md_figures" if with_images else "export_md"
        _mark_export_artifact(store, key, dest)
        return dest

    dest = out or (
        export_dir / ("paper_with_figures.docx" if with_images else "paper_final.docx")
    )
    dest = Path(dest)
    dest.parent.mkdir(parents=True, exist_ok=True)
    markdown_to_docx(body, dest, title=doc_title, base_dir=store.path)
    key = "export_docx_figures" if with_images else "export_docx"
    _mark_export_artifact(store, key, dest)
    return dest


def _mark_export_artifact(store: SessionStore, key: str, dest: Path) -> None:
    """Record the export in session.json.

    `--out` may point outside the session dir (e.g. ~/Desktop); relative_to()
    would raise there. Store relative paths when possible, absolute otherwise.
    """
    try:
        shown: str = str(dest.resolve().relative_to(store.path.resolve()))
    except ValueError:
        shown = str(dest.resolve())
    store.mark_artifact(key, shown)


def _article_body_keep_figures(text: str) -> str:
    """Strip revision metadata but keep the article and # Figures gallery."""
    figures_match = re.search(r"^#{1,2}[ \t]+Figures[ \t]*$", text, re.MULTILINE)
    if figures_match is None:
        return article_body(text)

    head = article_body(text[: figures_match.start()])
    # The gallery itself may be followed by metadata sections — drop those too.
    tail_body = article_body(text[figures_match.start() :])
    return (head.rstrip() + "\n\n" + tail_body).strip() + "\n"


def markdown_to_docx(
    markdown: str,
    dest: Path,
    *,
    title: str = "Document",
    base_dir: Path | None = None,
) -> Path:
    """Convert a subset of Markdown to .docx via python-docx."""
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        if shutil.which("pandoc"):
            return _pandoc_md_to_docx(markdown, dest)
        raise RuntimeError(
            "python-docx is required for Word export. Install with: "
            "uv pip install python-docx"
        ) from exc

    from council.media import ensure_raster_image

    base = base_dir or dest.parent
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(_strip_md_inline(h.group(2)), level=level)
            i += 1
            continue

        if re.match(r"^(-{3,}|\*{3,}|_{3,})\s*$", line.strip()):
            doc.add_paragraph("—" * 20)
            i += 1
            continue

        if re.match(r"^[-*+]\s+", line):
            while i < len(lines) and re.match(r"^[-*+]\s+", lines[i]):
                item = re.sub(r"^[-*+]\s+", "", lines[i])
                doc.add_paragraph(_strip_md_inline(item), style="List Bullet")
                i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                item = re.sub(r"^\d+\.\s+", "", lines[i])
                doc.add_paragraph(_strip_md_inline(item), style="List Number")
                i += 1
            continue

        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            p = doc.add_paragraph(_strip_md_inline(" ".join(quote_lines)))
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img:
            alt, src = img.group(1), img.group(2)
            img_path = Path(src)
            if not img_path.is_absolute():
                candidates = [
                    Path(src),
                    base / src,
                    base / "final" / src,
                    dest.parent / src,
                ]
                img_path = next((c for c in candidates if c.exists()), base / src)

            # Extract figure index from filename if possible
            idx = 1
            m_idx = re.search(r"figure_(\d+)", img_path.stem)
            if m_idx:
                idx = int(m_idx.group(1))

            raster = ensure_raster_image(
                img_path,
                title=alt or img_path.stem,
                caption=alt or "",
                index=idx,
            )
            if raster and raster.exists():
                try:
                    doc.add_picture(str(raster), width=Inches(6.0))
                    cap = doc.add_paragraph(alt or raster.name)
                    cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in cap.runs:
                        run.italic = True
                        run.font.size = Pt(9)
                except Exception:  # noqa: BLE001
                    doc.add_paragraph(f"[Image: {alt or src}]")
            else:
                doc.add_paragraph(f"[Image missing: {alt or src}]")
            i += 1
            continue

        para = line
        i += 1
        while (
            i < len(lines)
            and lines[i].strip()
            and not re.match(r"^(#{1,6}\s|[-*+]\s|\d+\.\s|>|```|!\[)", lines[i])
        ):
            para += " " + lines[i].strip()
            i += 1
        p = doc.add_paragraph()
        _add_runs_with_inline(p, para)

    doc.core_properties.title = title
    doc.save(str(dest))
    return dest


def _strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def _add_runs_with_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("["):
            label = re.match(r"\[([^\]]+)\]", token)
            run = paragraph.add_run(label.group(1) if label else token)
            run.underline = True
        else:
            paragraph.add_run(token)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _pandoc_md_to_docx(markdown: str, dest: Path) -> Path:
    import subprocess
    import tempfile

    with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(markdown)
        md_path = f.name
    try:
        subprocess.run(
            ["pandoc", md_path, "-o", str(dest)],
            check=True,
            capture_output=True,
            text=True,
        )
    finally:
        Path(md_path).unlink(missing_ok=True)
    return dest
